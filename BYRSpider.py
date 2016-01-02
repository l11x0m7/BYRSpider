# -*- coding: utf-8 -*-
from __future__ import division
import nltk, re, pprint,sys,os
import urllib, urllib2
import time, thread
from docx import Document
from docx.shared import Inches


def UrlToMobile(oriurl):
    pat = r'http://m.byr.cn/article/'
    now = oriurl.strip().split('/')
    if now[-1] == '':
        pat = pat + now[-3] +'/' + now[-2]
    else:
        pat = pat + now[-2] + '/' + now[-1]
    return pat

#----------- 主模块 -----------
class Spider_Model:

    def __init__(self, url):
        self.url = url
        self.page_num = 1
        self.pages = []
        self.img = dict()
        name = self.url.strip().split('/')
        name = name[-2] + name[-1]
        self.dirname = name
        if not os.path.exists(name):
            os.mkdir(name)

    def __GetPageNum(self):
        myUrl = self.url
        user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        headers = { 'User-Agent' : user_agent }
        req = urllib2.Request(myUrl, headers = headers)
        try:
            myResponse = urllib2.urlopen(req)
        except urllib2.URLError, e:
            print '[GetPageNum]:', e.code
        myPage = myResponse.read()
        match = re.findall(r'<a href=".*?\?p=(\d+)">尾页</a>', myPage, re.S)
        return int(match[0])


    def GetPage(self,page):
        myUrl = self.url + '?p=' + str(page)
        user_agent = 'Mozilla/4.0 (compatible; MSIE 5.5; Windows NT)'
        headers = { 'User-Agent' : user_agent }
        req = urllib2.Request(myUrl, headers = headers)
        try:
            myResponse = urllib2.urlopen(req)
        except urllib2.URLError, e:
            print '[GetPage]:', e.code
        myPage = myResponse.read()
        tag = False
        unicodePage = myPage.decode("utf-8")
        if u"精彩评论" in unicodePage:
            tag = True
        title = []
        floor = []
        people = []
        content = []
        time = []
        ip = []
        if int(page) == 1:
            title = re.findall(r'<li class="f">(.*?)</li>',unicodePage,re.S)

        floor = re.findall(r'<div class="nav hl"><div><a class="plant">(.*?)</a>', unicodePage, re.S)
        people = re.findall(r'<a href="/user/query/.*?">(.*?)</a>', unicodePage, re.S)
        time = re.findall(r'<a class="plant">(\d+-\d+-\d+ \d{2}:\d{2}:\d{2})</a>', unicodePage, re.S)
        ip = re.findall(r'<br />FROM (.*?)<span href', unicodePage, re.S)
        for i in range(len(ip)):
            myip = ip[i].decode()
            if '单击此查看原图' in myip:
                ip[i] = re.findall(r'<br />FROM (.*?)<br /><a target="_blank"', unicodePage, re.S)[0]


        # 匹配content
        content = re.findall(r'<li.*?<div class="sp">(.*?)FROM.*?</li>', unicodePage, re.S)
        if page=='1':
                print content
        # 过滤精彩评论
        if tag:
            content[1] = re.findall(r'<div class="sp">(.*?)<', content[1], re.S)[-1]
        for i in range(len(content)):
            # 消除各种可能的开头和结尾
            match = re.findall(r'(.*?)<br/><br/><a', content[i], re.S)
            if len(match)>0:
                content[i] = match[0]
            match = re.findall(r'(.*?)<br>--', content[i], re.S)
            if len(match)>0:
                content[i] = match[0]
            match = re.findall(r'<div class="sp">(.*)', content[i], re.S)
            if len(match)>0:
                content[i] = match[0]
            # 消除插入的图片并获取图片下载路径
            match = re.findall(r'(.*?)<a target="_blank".*?src="(.*?)".*?"/>(.*)', content[i], re.S)
            if page=='1' and i==0:
                print match
            real_content = ''
            while True:
            # if len(match)>0:
                if len(match)==0:
                    break
                for each in match:
                    real_content = real_content + each[0] + '\n'
                    self.img.setdefault(floor[i].decode(), [])
                    self.img[floor[i].decode()].append(each[1])
                    if r'<a target="_blank"' not in each[2]:
                        real_content += each[2]
                        match = ''
                        break
                    match = re.findall(r'(.*?)<a target="_blank".*?src="(.*?)".*?"/>(.*)', each[2], re.S)
            if len(real_content)>0:
                content[i] = real_content
            # 消除表情
            match = re.findall(r'(.*?)<img src=.*?/>(.*)', content[i], re.S)
            real_content = ''
            # if len(match)>0:
            #     for each in match:
            #         content[i] = each[0] + '\n' + each[1]
            while True:
                if len(match)==0:
                    break
                for each in match:
                    real_content = real_content + each[0]
                    if r'img src=' not in each[1]:
                        real_content += each[1]
                        match = ''
                        break
                    match = re.findall(r'(.*?)<img src=.*?/>(.*)', each[1], re.S)
            if len(real_content)>0:
                content[i] = real_content


            # if len(match)>0:
            content[i] = content[i].replace("<br/>", '\n')
            content[i] = content[i].replace("<br />", '\n')
            content[i] = content[i].replace('&nbsp;', ' ')
        try:
            result = []
            if tag:
                result.append(title)
            for i in range(len(content)):
                temp = [floor[i], people[i], time[i], ip[i], content[i]]
                result.append(temp)
            return result
        except Exception, e:
            print '错误，未对齐！'
            print 'floor：%d' % len(floor)
            print 'content：%d' % len(content)
            print 'people：%d' % len(people)
            print 'time：%d' % len(time)
            print 'ip：%d' % len(ip)

    def getImg(self, pat):
        name = pat.strip().split('/')
        name = name[-3] + name[-2]
        imgurl = r'http://m.byr.cn' + pat
        urllib.urlretrieve(imgurl, self.dirname + '/' + name + '.jpg')
        return name + '.jpg'

    # 用于加载一个帖子内容
    def LoadWholePages(self):
        self.pages = []
        self.page_num = self.__GetPageNum()
        for page in range(1,self.page_num+1):
            try:
                # 获取页面
                myPage = self.GetPage(str(page))
                self.pages.append(myPage)
            except:
                print '第%d页无法连接到BYR论坛！' % page
            time.sleep(1)

    def ShowTiezi(self):
        if len(self.pages)==0:
            print 'Error!'
        for i in range(self.page_num):
            now_page = self.pages[i]
            for a in now_page:
                for b in a:
                    print b.decode()
                print '---------------------------------'

    def WriteTiezi(self):
        if len(self.pages)==0:
            print 'Error!'
        document = Document()
        # style = document.StyleSheet
        # style.Fonts.append(Font("\\'cb\\'ce\\'cc\\'e5", 'modern', 134, 30))
        # section = Section()
        # document.Sections.append(section)
        # tps = TextPS(font=getattr(style.Fonts, "\\'cb\\'ce\\'cc\\'e5"))
        for i in range(self.page_num):
                now_page = self.pages[i]
                for a in now_page:
                    for b in a:
                        document.add_paragraph(b.decode())
                    if len(a)>1:
                        # pic_num = len(self.img[a[0]])
                        if self.img.has_key(a[0]):
                            for k in self.img[a[0].decode()]:
                                pic_name = self.getImg(k)
                                document.add_picture(self.dirname + '/' + pic_name)
                    document.add_paragraph('---------------------------------')
        name = self.url.strip().split('/')
        name = name[-2] + name[-1]
        document.save(self.dirname + '/' + name + '.docx')
        print "Success to dump into " + name + '.docx'




if __name__ == '__main__':
    #----------- 程序的入口处 -----------
    print """
    ---------------------------------------
       程序：BYR论坛爬虫
       版本：0.1
       作者：lxm
       日期：2016-01-01
       语言：Python 2.7
       功能：爬取北邮人论坛帖子
    ---------------------------------------
    """
    reload(sys)
    sys.setdefaultencoding('utf-8')
    # headers = {
    #     'User-Agent':'Mozilla/5.0 (Windows; U; Windows NT 6.1; en-US; rv:1.9.1.6) Gecko/20091201 Firefox/3.5.6'
    # }
    url = 'http://bbs.byr.cn/#!article/Java/46891'
    data = urllib.urlencode({})
    newurl = UrlToMobile(url)
    spider = Spider_Model(newurl)
    # req = urllib2.Request(newurl, data, headers)
    # res = urllib2.urlopen(req)
    spider.LoadWholePages()
    spider.WriteTiezi()
